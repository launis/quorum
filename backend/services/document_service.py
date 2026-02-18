"""Document Service for file processing."""

import io
import logging
import os
from typing import Any, Dict, Tuple, Union

import docx
import fitz  # PyMuPDF
from fastapi import status
from fastapi.concurrency import run_in_threadpool

from backend.exceptions import AppException, ErrorCodes, FatalInterruption
from backend.services.chat_log_parser import ChatLogParser
from backend.services.file_driver import FileDriver
from backend.services.knowledge_base_parser import KnowledgeBaseParser
from backend.models.domain.knowledge import KnowledgeBaseDocument

logger = logging.getLogger(__name__)


class DocumentService:
    """Unified service for handling document ingestion, processing, and archiving.

    Supports:
    - Evidence files (PDF/DOCX -> Text) for WorkflowEngine
    - Knowledge Base files (DOCX/MD -> Structured JSON) for KnowledgeBaseService.

    Architecture:
    - Uses 'run_in_threadpool' for CPU-bound tasks (OCR/Extraction).
    - Uses FileDriver for file persistence.
    """

    def __init__(self, storage_client: FileDriver):
        """Initializes the service.

        Args:
            storage_client (FileDriver): The storage backend.

        """
        self.storage_client = storage_client

    async def process_evidence_files(self, execution_id: str, files: Dict[str, Tuple[str, bytes]]) -> Dict[str, str]:
        """Archives evidence files to storage and extracts text for workflow execution.

        Handles PDF and DOCX formats automatically.

        Args:
            execution_id (str): UUID of the execution context.
            files (Dict[str, Tuple[str, bytes]]): Map of {input_key: (filename, content)}.

        Returns:
            Dict[str, str]: Map of {input_key: extracted_text_content}.

        Raises:
            AppException: If file processing fails critically (Fail Fast).
        """
        # FAIL FAST: Empty Input
        if not files:
            # Not necessarily an error if no files were uploaded, but if the step relied on it...
            # The calling step usually checks if input is missing.
            # Here we just return empty dict if empty, but if files are provided, we process strictly.
            return {}

        extracted_data: Dict[str, str] = {}

        for input_key, (filename, file_bytes) in files.items():
            try:
                # 1. Archive to Storage (IO-bound)
                relative_path = f"{execution_id}/{filename}"
                saved_path = await self.storage_client.save(relative_path, file_bytes)

                # 2. Extract Text (CPU-bound)
                lower_name = filename.lower()
                text = ""

                if lower_name.endswith(".pdf"):
                    text = await run_in_threadpool(self._extract_text_from_pdf, file_bytes)
                elif lower_name.endswith(".docx"):
                    text = await run_in_threadpool(self._extract_text_from_docx, file_bytes)
                else:
                    # Treat as text file
                    text = file_bytes.decode("utf-8", errors="ignore")

                # --- NEW: Parse Chat Logs ---
                # Attempt to identify and label speakers (User/AI) to assist the Profiler.
                # ChatLogParser now has FAIL FAST checks.
                text = ChatLogParser.parse(text)
                # ----------------------------

                extracted_data[input_key] = text

                logger.info(
                    f"[DocumentService] Evidence {filename} processed. Extracted {len(text)} chars. "
                    f"Storage: {saved_path}"
                )

            except Exception as e:
                # Catch-all to wrap in AppException (RFC 7807)
                logger.error(f"[DocumentService] Failed to ingest evidence {filename} ({input_key}): {e}")
                
                # If it's already an AppException, re-raise
                if isinstance(e, AppException):
                    raise e
                
                raise AppException(
                    message=f"Failed to ingest evidence {filename}: {str(e)}",
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    details={
                        "error_code": ErrorCodes.DOCUMENT_PROCESSING_FAILED,
                        "filename": filename,
                        "original_error": str(e)
                    }
                ) from e

        return extracted_data

    async def process_knowledge_base_file(self, content: bytes, filename: str, job_id: str) -> KnowledgeBaseDocument:
        """Archives Knowledge Base file and parses it into concepts/references.

        Supports both DOCX and Markdown formats.

        Args:
            content (bytes): File content.
            filename (str): Name of the file.
            job_id (str): Background task ID for storage organization.

        Returns:
            KnowledgeBaseDocument: Structured data for ingestion.

        Raises:
            AppException: If file type is unsupported or parsing fails.
        """
        # FAIL FAST: Empty Content
        if not content:
             raise AppException(
                message="Knowledge Base file content is empty.",
                status_code=status.HTTP_400_BAD_REQUEST,
                details={"error_code": ErrorCodes.INVALID_FILE_FORMAT, "filename": filename}
            )

        is_docx = filename.lower().endswith(".docx")
        is_md = filename.lower().endswith(".md")

        if not (is_docx or is_md):
            raise AppException(
                message=f"Knowledge Base must be a DOCX or MD file. Got: {filename}",
                status_code=status.HTTP_400_BAD_REQUEST,
                details={"error_code": ErrorCodes.INVALID_FILE_FORMAT, "filename": filename}
            )

        try:
            # 1. Archive
            relative_path = f"knowledge_base/{job_id}/{filename}"
            await self.storage_client.save(relative_path, content)

            # 2. Parse (CPU-bound)
            parsed_doc = None
            if is_docx:
                # Wrap bytes in stream for parser
                file_stream = io.BytesIO(content)
                # Pass filename as kwarg or arg? run_in_threadpool takes *args.
                # Signature: parse_docx(file_input: Any, filename: str = ...)
                parsed_doc = await run_in_threadpool(KnowledgeBaseParser.parse_docx, file_stream, filename)
            else:
                # Must be MD due to validation above
                # Wrap bytes in stream for parser (parse_md handles extraction)
                file_stream = io.BytesIO(content)
                parsed_doc = await run_in_threadpool(KnowledgeBaseParser.parse_md, file_stream, filename)

            logger.info(
                f"[DocumentService] KB {filename} processed. Found {len(parsed_doc.concepts)} concepts."
            )
            return parsed_doc

        except Exception as e:
            logger.error(f"[DocumentService] KB processing failed for {filename}: {e}")
            if isinstance(e, AppException):
                raise e
            raise AppException(
                message=f"Knowledge Base processing failed: {str(e)}",
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                details={"error_code": ErrorCodes.KNOWLEDGE_INGESTION_FAILED, "original_error": str(e)}
            ) from e

    # --- Internal Text Extraction Helpers (Migrated from DocumentProcessor) ---

    def extract_text(self, input_data: Union[str, bytes]) -> str:
        """Unified text extraction method (Public API).

        Routes to PDF or DOCX extractors based on content or filename.

        Args:
            input_data (Union[str, bytes]): File path or file content bytes.
        
        Returns:
            str: Extracted text.

        Raises:
            AppException: If extraction fails or file not found.
        """
        try:
            if isinstance(input_data, str):
                # FAIL FAST: Check file existence
                if not os.path.exists(input_data):
                    raise AppException(
                        message=f"File not found: {input_data}",
                        status_code=status.HTTP_404_NOT_FOUND,
                        details={"error_code": ErrorCodes.FILE_NOT_FOUND, "path": input_data}
                    )

                # Path based dispatch
                lower = input_data.lower()
                if lower.endswith(".pdf"):
                    return self._extract_text_from_pdf(input_data)
                elif lower.endswith(".docx"):
                    return self._extract_text_from_docx(input_data)
                else:
                     # Try simple read for text files
                    with open(input_data, encoding="utf-8") as f:
                        return f.read()

            elif isinstance(input_data, bytes):
                # In-memory dispatch using magic numbers or caller context?
                # This seems risky for strict Fail Fast if we guess wrong. 
                # But for now, we assume if bytes are passed, we might need a type hint arg in future.
                # Current usage usually involves specific methods (process_evidence_files).
                # extract_text is a helper.
                pass

            # Fail Fast if we reach here
            raise AppException(
                message="Cannot determine file type or process input data.",
                status_code=status.HTTP_400_BAD_REQUEST,
                details={"error_code": ErrorCodes.INVALID_FILE_FORMAT}
            )

        except Exception as e:
            if isinstance(e, AppException):
                raise e
            raise AppException(
                message=f"Text extraction failed: {str(e)}",
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                details={"error_code": ErrorCodes.DOCUMENT_PROCESSING_FAILED, "original_error": str(e)}
            ) from e

    @staticmethod
    def _extract_text_from_pdf(input_data: Union[str, bytes]) -> str:
        """Extracts plain text from a PDF file using PyMuPDF (fitz)."""
        try:
            doc = None
            if isinstance(input_data, str):
                if not os.path.exists(input_data):
                    raise FileNotFoundError(f"File not found: {input_data}")
                doc = fitz.open(input_data)
            elif isinstance(input_data, bytes):
                doc = fitz.open(stream=input_data, filetype="pdf")
            else:
                raise ValueError("Input must be a file path (str) or bytes.")

            text = ""
            for page in doc:
                text += page.get_text()

            return text.strip()
        except Exception as e:
            raise AppException(
                message=f"Failed to extract PDF text: {str(e)}",
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                details={"error_code": ErrorCodes.DOCUMENT_PROCESSING_FAILED, "format": "pdf"}
            ) from e

    @staticmethod
    def _extract_text_from_docx(input_data: Union[str, bytes]) -> str:
        """Extracts plain text from a DOCX file using python-docx."""
        try:
            doc = None
            if isinstance(input_data, str):
                if not os.path.exists(input_data):
                    raise FileNotFoundError(f"File not found: {input_data}")
                doc = docx.Document(input_data)
            elif isinstance(input_data, bytes):
                file_stream = io.BytesIO(input_data)
                doc = docx.Document(file_stream)
            else:
                raise ValueError("Input must be a file path (str) or bytes.")

            text = []
            for para in doc.paragraphs:
                text.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            text.append(para.text)

            return "\n".join(text).strip()
        except Exception as e:
             raise AppException(
                message=f"Failed to extract DOCX text: {str(e)}",
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                details={"error_code": ErrorCodes.DOCUMENT_PROCESSING_FAILED, "format": "docx"}
            ) from e
