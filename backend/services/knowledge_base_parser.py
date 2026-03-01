"""Parser service for extracting structured knowledge from documents."""

import logging
import re
import uuid
from datetime import datetime, timezone
from typing import Any

import docx
from fastapi import status

from backend.exceptions import AppException, ErrorCodes
from backend.models.domain.knowledge import Claim, Concept, DocumentChunk, KnowledgeBaseDocument, Reference

logger = logging.getLogger(__name__)


class KnowledgeBaseParser:
    """Parse unstructured documents (DOCX, Markdown) into structured Knowledge Base entries.

    Extract Concepts, References, and Claims using heuristic detection and ReGex patterns.
    Support Multilingual Headers (Fi, En).
    """

    BIBLIOGRAPHY_HEADERS = {
        "lähdeluettelo",
        "lähteet",
        "bibliography",
        "references",
        "works cited",
    }

    STRUCTURAL_KEYWORDS = {
        "lähdeluettelo",
        "lähteet",
        "abstrakti",
        "tiivistelmä",
        "analyysi",
        "johdanto",
        "bibliography",
        "references",
        "abstract",
        "summary",
        "analysis",
        "introduction",
    }

    COMPARE_MARKERS = ["vrt.", "cf."]

    AND_MARKERS = [" and ", " ja ", " & "]

    @staticmethod
    def extract_claims_from_text(text: str) -> list[Claim]:
        """Extract sentences containing citations (claims) from text using ReGex.

        Return strictly typed Claim objects.
        """
        claims: list[Claim] = []
        if not text:
            return claims

        # Regex definitions
        link_pattern = re.compile(r"\[([^\]]+)\]\(#([a-zA-Z0-9_-]+)\)")
        text_cit_pattern = re.compile(
            r"(\((?:(?:vrt\.|cf\.)\s*)?(?:[\w&,-]+\s+)+(?:et\s+al\.|ym\.)?\s*,?\s*\d{4}[a-z]?\))", re.IGNORECASE
        )

        replacements: dict[str, str] = {}

        def mask_match(match):
            key = f"__CIT_MASK_{len(replacements)}__"
            replacements[key] = match.group(0)
            return key

        masked_text = link_pattern.sub(mask_match, text)
        masked_text = text_cit_pattern.sub(mask_match, masked_text)

        sentences = re.split(r"(?<=[.!?])\s+", masked_text)

        for sent in sentences:
            sent = sent.strip()
            if len(sent) < 10:
                continue

            found_masks = re.findall(r"__CIT_MASK_\d+__", sent)

            if found_masks:
                citation_labels = []
                citation_keys = []
                matches_text_citation = False

                clean_claim = sent

                for mask in found_masks:
                    original = replacements[mask]

                    # Analyze the original citation
                    # Determine if it's Link or Text
                    # Simple heuristic: starts with '[' is link, '(' is text
                    if original.startswith("["):
                        # Link
                        m = link_pattern.match(original)
                        if m:
                            citation_labels.append(m.group(1))
                            citation_keys.append(m.group(2))
                    elif original.startswith("("):
                        # Text Citation
                        matches_text_citation = True
                        # Parse internals
                        # Remove parens
                        inner = original.strip("()")
                        # Remove markers
                        for marker in KnowledgeBaseParser.COMPARE_MARKERS:
                            if inner.lower().startswith(marker):
                                inner = inner[len(marker) :].strip()

                        # Deduplicate labels
                        if inner not in citation_labels:
                            citation_labels.append(inner)

                    # Remove mask from Clean Claim
                    clean_claim = clean_claim.replace(mask, "")

                # Cleanup Clean Claim
                clean_claim = re.sub(r"\s+", " ", clean_claim)
                clean_claim = re.sub(r"\s+\.", ".", clean_claim)
                clean_claim = re.sub(r"\s+,", ",", clean_claim)
                clean_claim = re.sub(r"\(\)", "", clean_claim)  # Should be handled by mask removal but good safety
                clean_claim = clean_claim.strip()

                # Filter junk
                if len(re.findall(r"[a-zA-ZäöåÄÖÅ0-9]", clean_claim)) < 3:
                    continue
                if len(clean_claim) < 5:
                    continue

                claims.append(
                    Claim(
                        claim_text=clean_claim,
                        citation_keys=citation_keys,
                        citation_text="; ".join(citation_labels),
                        original_markdown=None,  # Resolved later
                        matches_text_citation=matches_text_citation,
                    )
                )
        return claims

    @staticmethod
    def parse_docx(file_input: Any, filename: str = "unknown.docx") -> KnowledgeBaseDocument:
        """Parse DOCX document into structured knowledge.

        Iterate through paragraphs to distinguish between Concepts (Headers + Text) and Bibliography.

        Args:
            file_input (Any): File path (str) or file-like object (stream).
            filename (str): Original filename for metadata.

        Returns:
            KnowledgeBaseDocument: Strict Pydantic model.

        Raises:
            AppException: If document cannot be opened (PARSING_FAILED).

        """
        logger.info(f"[KBParser] Parsing input (Type: {type(file_input)})")
        try:
            doc = docx.Document(file_input)
        except Exception as e:
            logger.error(f"[KBParser] Failed to open document: {e}")
            raise AppException(
                message=f"Failed to open DOCX document: {str(e)}",
                status_code=status.HTTP_400_BAD_REQUEST,
                details={"error_code": ErrorCodes.INVALID_FILE_FORMAT, "original_error": str(e)},
            ) from e

        # Data containers
        concepts: list[Concept] = []
        references: list[Reference] = []
        claims: list[Claim] = []

        # We also collect raw chunks for the document model
        chunks: list[DocumentChunk] = []
        total_tokens = 0  # Estimate

        # Regex for DOI
        doi_pattern = re.compile(r"\b(10.\d{4,9}/[-._;()/:A-Z0-9]+)\b", re.IGNORECASE)

        current_concept = None
        current_definition: list[str] = []

        in_bibliography = False

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Heuristics
            style_name = ""
            if para.style and hasattr(para.style, "name"):
                style_name = para.style.name.lower()

            # Detect Bibliography Section (Multilingual)
            if any(t in text.lower() for t in KnowledgeBaseParser.BIBLIOGRAPHY_HEADERS):
                # Loose check: Heading style OR bold OR distinct short text
                is_header = "heading" in style_name or len(text) < 60

                if is_header:
                    in_bibliography = True
                    # Close previous concept if open
                    if current_concept:
                        def_text = "\n".join(current_definition)
                        concepts.append(Concept(term=current_concept, definition=def_text))
                        # Extract Claims
                        section_claims = KnowledgeBaseParser.extract_claims_from_text(def_text)
                        for c in section_claims:
                            # Update context (since Claim is frozen, we might need to recreate or use dict first?
                            # Claim is Pydantic. We can use model_copy if needed, but easier to construct directly.
                            # extract_claims now returns Claim objects.
                            # We need to set concept_context. Models are frozen=True.
                            # So we must reconstitute.
                            claims.append(
                                Claim(
                                    claim_text=c.claim_text,
                                    citation_keys=c.citation_keys,
                                    citation_text=c.citation_text,
                                    original_markdown=c.original_markdown,
                                    matches_text_citation=c.matches_text_citation,
                                    concept_context=current_concept,
                                )
                            )

                        current_concept = None
                        current_definition = []

                    logger.info(f"[KBParser] Bibliography section start detected at: '{text}'")
                    continue

            if in_bibliography:
                citation = text
                doi_match = doi_pattern.search(citation)
                doi_link = None
                if doi_match:
                    doi_raw = doi_match.group(1)
                    doi_link = f"https://doi.org/{doi_raw}"

                references.append(
                    Reference(
                        citation=citation,
                        short_citation=KnowledgeBaseParser.extract_short_citation(citation),
                        doi_link=doi_link,
                    )
                )
            else:
                if "heading" in style_name and "1" not in style_name:
                    if current_concept:
                        def_text = "\n".join(current_definition)
                        if def_text.strip():
                            concepts.append(Concept(term=current_concept, definition=def_text))
                            # Extract Claims
                            section_claims = KnowledgeBaseParser.extract_claims_from_text(def_text)
                        else:
                            section_claims = []  # No text, no claims
                        for c in section_claims:
                            claims.append(
                                Claim(
                                    claim_text=c.claim_text,
                                    citation_keys=c.citation_keys,
                                    citation_text=c.citation_text,
                                    original_markdown=c.original_markdown,
                                    matches_text_citation=c.matches_text_citation,
                                    concept_context=current_concept,
                                )
                            )

                    current_concept = text
                    current_definition = []
                elif "heading" in style_name or (len(text) < 50 and text.isupper()):
                    # GENERALIZED LOGIC: Use Heading Levels

                    is_structural = False

                    # 1. Check Style Level
                    if "heading 1" in style_name or "title" in style_name:
                        is_structural = True

                    # 2. Numbering Heuristic (Universal)
                    match_num = re.match(r"^(\d+(\.\d+)*)\.?\s+", text)
                    if match_num:
                        numbering = match_num.group(1)  # "1" or "1.2"
                        dot_count = numbering.count(".")
                        if dot_count == 0:
                            is_structural = True

                    # 3. Fallback: Structural Keywords (Multilingual)
                    if any(
                        k == text.lower().strip() or (text.lower().startswith(k) and len(text) < 30)
                        for k in KnowledgeBaseParser.STRUCTURAL_KEYWORDS
                    ):
                        is_structural = True

                    if current_concept:
                        def_text = "\n".join(current_definition)
                        if def_text.strip():
                            concepts.append(Concept(term=current_concept, definition=def_text))
                            # Extract Claims
                            section_claims = KnowledgeBaseParser.extract_claims_from_text(def_text)
                        else:
                            section_claims = []

                        for c in section_claims:
                            claims.append(
                                Claim(
                                    claim_text=c.claim_text,
                                    citation_keys=c.citation_keys,
                                    citation_text=c.citation_text,
                                    original_markdown=c.original_markdown,
                                    matches_text_citation=c.matches_text_citation,
                                    concept_context=current_concept,
                                )
                            )

                    if is_structural:
                        logger.info(f"[KBParser] Skipping structural/H1 header: '{text}'")
                        current_concept = None
                    else:
                        current_concept = text

                    current_definition = []
                else:
                    if current_concept:
                        current_definition.append(text)

        if current_concept:
            def_text = "\n".join(current_definition)
            if def_text.strip():
                concepts.append(Concept(term=current_concept, definition=def_text))
                section_claims = KnowledgeBaseParser.extract_claims_from_text(def_text)
            else:
                section_claims = []
            for c in section_claims:
                claims.append(
                    Claim(
                        claim_text=c.claim_text,
                        citation_keys=c.citation_keys,
                        citation_text=c.citation_text,
                        original_markdown=c.original_markdown,
                        matches_text_citation=c.matches_text_citation,
                        concept_context=current_concept,
                    )
                )

        claims = KnowledgeBaseParser._resolve_claims(claims, references)

        # Build chunks from concepts for now (one chunk per concept)
        for i, concept in enumerate(concepts):
            chunks.append(
                DocumentChunk(
                    chunk_id=f"chunk-{i}",
                    content=f"{concept.term}\n\n{concept.definition}",
                    page_number=None,
                    metadata={"term": concept.term},
                )
            )
            total_tokens += len(concept.definition.split())  # Rough estimate

        logger.info(
            f"[KBParser] Extracted {len(concepts)} concepts, "
            f"{len(references)} references, and {len(claims)} claims from DOCX."
        )

        return KnowledgeBaseDocument(
            document_id=str(uuid.uuid4()),
            filename=filename,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            total_tokens=total_tokens,
            chunks=chunks,
            concepts=concepts,
            references=references,
            claims=claims,
            parsed_at=datetime.now(timezone.utc),
            metadata={"source": "docx_upload"},
        )

    @staticmethod
    def clean_text(text: str) -> str:
        """Normalize text by removing invisible Word artifacts (non-breaking spaces, dashes)."""
        if not text:
            return ""
        text = text.replace("\xa0", " ").replace("–", "-").replace("—", "-")
        text = re.sub(r"\s+", " ", text).strip()
        return text

    @staticmethod
    def extract_short_citation(full_entry: str) -> str | None:
        """Extract concise 'Author Year' or 'Author & Author Year' label from a full bibliographic entry.

        Supported formats:
        - "Acemoglu, D. & Restrepo, P. 2018:"
        - "Acemoglu (2018)."
        - "Acemoglu, D. 2018."

        Args:
            full_entry (str): The full bibliography line.

        Returns:
            Optional[str]: Short citation string or None.

        """
        if not full_entry:
            return None

        full_entry = KnowledgeBaseParser.clean_text(full_entry)

        # Regex strategies
        patterns = [
            # 1. Standard: "Name. 2023:" or "Name. 2023."
            r"^(.+?)\.\s*(\d{4}[a-z]?)[.:]",
            # 2. Parentheses: "Name (2023)."
            r"^(.+?)\s*\((\d{4}[a-z]?)\)",
            # 3. No dot: "Name 2023" (rare but possible in headers)
            r"^(.+?)\s+(\d{4}[a-z]?)",
        ]

        match = None
        for p in patterns:
            m = re.match(p, full_entry)
            if m:
                match = m
                break

        if match:
            authors_part = match.group(1)
            year = match.group(2)

            # Remove (toim.) / (eds.)
            authors_part = re.sub(r"\s*\(.*?\)", "", authors_part)

            # Extract surnames
            authors = []

            # Normalize separators to & (Multilingual)
            for sep in KnowledgeBaseParser.AND_MARKERS:
                authors_part = authors_part.replace(sep, " & ")

            parts = authors_part.split("&")
            for part in parts:
                # Split by comma to get surname (assumes "Surname, Firstname")
                # If just "Surname", split returns [Surname]
                names = re.split(r"[,]", part)
                surname = names[0].strip()
                # Simple heuristic: Surname should have letters
                if surname and any(c.isalpha() for c in surname):
                    authors.append(surname)

            # Construct short citation
            if not authors:
                # Fallback: use first word of authors_part if parsing failed but regex matched
                authors = [authors_part.split()[0]]

            short_authors = " & ".join(authors)
            if len(authors) > 2:
                # Multilingual "et al" - defaulting to "ym." if mostly Finnish context, or "et al."
                # Ideally detection? Defaulting to "et al." as strict academic standard, or "ym." for consistency with regex.
                # Let's use "et al." as international standard unless detected otherwise.
                # Actually, previous code used "ym." hardcoded!
                short_authors = f"{authors[0]} et al."

            return f"{short_authors} {year}"
        return None

    @staticmethod
    def parse_md(file_input: Any, filename: str = "unknown.md") -> KnowledgeBaseDocument:
        """Parse Markdown content into structured knowledge.

        Support for:
        - Headers (# Term) -> Concepts
        - Lists (- Ref) in Bibliography section -> References
        - Anchor IDs ({#id}) linking

        Args:
            file_input (Any): File path (str) or file-like object (stream).
            filename (str): Original filename.

        Returns:
             KnowledgeBaseDocument: Structured KB model.

        Raises:
            AppException: If parsing fails (PARSING_FAILED).
        """
        logger.info(f"[KBParser] Parsing MD input (Type: {type(file_input)})")

        content_str = ""
        try:
            if isinstance(file_input, str):
                with open(file_input, encoding="utf-8") as f:
                    content_str = f.read()
            else:
                # Assume stream/bytes
                if hasattr(file_input, "read"):
                    content = file_input.read()
                    if isinstance(content, bytes):
                        content_str = content.decode("utf-8")
                    else:
                        content_str = content
                elif isinstance(file_input, bytes):
                    content_str = file_input.decode("utf-8")
        except Exception as e:
            logger.error(f"[KBParser] Failed to read MD content: {e}")
            raise AppException(
                message=f"Failed to read Markdown content: {str(e)}",
                status_code=status.HTTP_400_BAD_REQUEST,
                details={"error_code": ErrorCodes.INVALID_FILE_FORMAT, "original_error": str(e)},
            ) from e

        lines = content_str.splitlines()

        # Data containers
        concepts: list[Concept] = []
        references: list[Reference] = []
        claims: list[Claim] = []
        chunks: list[DocumentChunk] = []
        total_tokens = 0

        # Regex for DOI
        doi_pattern = re.compile(r"\b(10.\d{4,9}/[-._;()/:A-Z0-9]+)\b", re.IGNORECASE)

        # Regex for Bibliography Header (Multilingual)
        # Construct pattern from BIBLIOGRAPHY_HEADERS
        headers_pattern_str = "|".join(re.escape(h) for h in KnowledgeBaseParser.BIBLIOGRAPHY_HEADERS)

        # Allow HTML tags (like <a id="...">) between #'s and the text
        bib_header_pattern = re.compile(r"^#+\s*(?:<[^>]+>\s*)*(" + headers_pattern_str + r")", re.IGNORECASE)

        # Regex for Concept Header
        concept_header_pattern = re.compile(r"^(#+)\s*(.+)")

        # Regex for Anchor ID in Bibliography: {#id} or <a id="id">
        anchor_pattern = re.compile(r'(?:\{\#([a-zA-Z0-9_-]+)\}|<a\s+id="([a-zA-Z0-9_-]+)">)')

        current_concept = None
        current_definition: list[str] = []

        in_bibliography = False

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Check for Bibliography Header
            if bib_header_pattern.match(line):
                in_bibliography = True
                # Close previous concept
                if current_concept:
                    def_text = "\n".join(current_definition)
                    concepts.append(Concept(term=current_concept, definition=def_text))
                    # Extract Claims from this section
                    section_claims = KnowledgeBaseParser.extract_claims_from_text(def_text)
                    for c in section_claims:
                        claims.append(
                            Claim(
                                claim_text=c.claim_text,
                                citation_keys=c.citation_keys,
                                citation_text=c.citation_text,
                                original_markdown=c.original_markdown,
                                matches_text_citation=c.matches_text_citation,
                                concept_context=current_concept,
                            )
                        )

                    current_concept = None
                    current_definition = []
                logger.info("[KBParser] Bibliography section detected (MD).")
                continue

            if in_bibliography:
                # Capture bullet points as references
                if line.startswith(("*", "-")):
                    citation = line.lstrip("*- ").strip()

                    # 1. Extract Anchor ID if present
                    anchor_id = None
                    anchor_match = anchor_pattern.search(citation)
                    if anchor_match:
                        # Group 1 is {#id}, Group 2 is <a id>
                        anchor_id = anchor_match.group(1) or anchor_match.group(2)

                        # Remove the anchor part from citation text
                        citation = anchor_pattern.sub("", citation).strip()
                        # Also remove [] if left over from []{#id}
                        citation = citation.replace("[]", "").strip()

                    # 2. Find DOI
                    doi_match = doi_pattern.search(citation)
                    doi_link = None
                    if doi_match:
                        doi_raw = doi_match.group(1)
                        doi_link = f"https://doi.org/{doi_raw}"

                    # 3. Clean
                    clean_citation = re.sub(r"<[^>]+>", "", citation)

                    references.append(
                        Reference(
                            citation=clean_citation,
                            short_citation=KnowledgeBaseParser.extract_short_citation(clean_citation),
                            doi_link=doi_link,
                            anchor_id=anchor_id,
                        )
                    )
                    # Continuation of previous reference? Or just ignore?
                    # For now, ignore non-bullet lines in bibliography unless strictly needed.
            else:
                # Concepts Processing
                header_match = concept_header_pattern.match(line)
                if header_match:
                    # H1, H2, H3...
                    # Generally treat header text as concept term
                    term = header_match.group(2).strip()

                    # Close previous
                    if current_concept:
                        def_text = "\n".join(current_definition)
                        concepts.append(Concept(term=current_concept, definition=def_text))
                        # Extract Claims from this section
                        section_claims = KnowledgeBaseParser.extract_claims_from_text(def_text)
                        for c in section_claims:
                            claims.append(
                                Claim(
                                    claim_text=c.claim_text,
                                    citation_keys=c.citation_keys,
                                    citation_text=c.citation_text,
                                    original_markdown=c.original_markdown,
                                    matches_text_citation=c.matches_text_citation,
                                    concept_context=current_concept,
                                )
                            )

                    # Remove anchor tags from term if present (e.g. <a id="foo"></a>Term)
                    term_clean = re.sub(r"<[^>]+>", "", term).strip()

                    current_concept = term_clean
                    current_definition = []
                else:
                    if current_concept:
                        current_definition.append(line)

        if current_concept:
            def_text = "\n".join(current_definition)
            concepts.append(Concept(term=current_concept, definition=def_text))
            section_claims = KnowledgeBaseParser.extract_claims_from_text(def_text)
            for c in section_claims:
                claims.append(
                    Claim(
                        claim_text=c.claim_text,
                        citation_keys=c.citation_keys,
                        citation_text=c.citation_text,
                        original_markdown=c.original_markdown,
                        matches_text_citation=c.matches_text_citation,
                        concept_context=current_concept,
                    )
                )

        claims = KnowledgeBaseParser._resolve_claims(claims, references)

        # Build chunks
        for i, concept in enumerate(concepts):
            chunks.append(
                DocumentChunk(
                    chunk_id=f"chunk-{i}",
                    content=f"{concept.term}\n\n{concept.definition}",
                    page_number=None,
                    metadata={"term": concept.term},
                )
            )
            total_tokens += len(concept.definition.split())

        logger.info(
            f"[KBParser] Extracted {len(concepts)} concepts, "
            f"{len(references)} references, and {len(claims)} claims from MD."
        )

        return KnowledgeBaseDocument(
            document_id=str(uuid.uuid4()),
            filename=filename,
            content_type="text/markdown",
            total_tokens=total_tokens,
            chunks=chunks,
            concepts=concepts,
            references=references,
            claims=claims,
            parsed_at=datetime.now(timezone.utc),
            metadata={"source": "markdown_upload"},
        )

    @staticmethod
    def _resolve_claims(claims: list[Claim], references: list[Reference]) -> list[Claim]:
        """Resolve textual claims to their full bibliographic references safely.
        
        Returns a new array with functionally recreated objects to preserve mutability constraints.
        """
        id_map = {}
        short_map = {}

        for r in references:
            if r.anchor_id:
                id_map[r.anchor_id] = r.citation

            sc = r.short_citation
            if sc:
                short_map[sc.lower()] = r.citation

        resolved_claims = []

        for c in claims:
            full_ref = None

            if c.citation_keys:
                for key in c.citation_keys:
                    if key in id_map:
                        full_ref = id_map[key]
                        break

            if not full_ref and c.citation_text:
                lbl = c.citation_text.split(";")[0].strip().lower()
                if lbl in short_map:
                    full_ref = short_map[lbl]
                else:
                    for s_key, s_val in short_map.items():
                        if lbl in s_key or s_key in lbl:
                            full_ref = s_val
                            break

            if full_ref:
                resolved_claims.append(c.model_copy(update={"original_markdown": full_ref}))
            else:
                resolved_claims.append(c)

        return resolved_claims
