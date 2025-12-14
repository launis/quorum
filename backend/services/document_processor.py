import fitz  # PyMuPDF
import os
import logging
import io
import docx
from typing import Dict, Any, Union
from backend.core.component import BaseComponent

logger = logging.getLogger(__name__)

class DocumentProcessor(BaseComponent):
    """
    Handles processing of documents (PDF, DOCX) to extract text.
    """

    def execute(self, file_path: str, **kwargs) -> Dict[str, Any]:
        """
        Executes the extraction based on file extension.
        """
        logger.info(f"Executing DocumentProcessor for file: {file_path}")
        if file_path.lower().endswith(".pdf"):
            text = self.extract_text_from_pdf(file_path)
        elif file_path.lower().endswith(".docx"):
            text = self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_path}")
            
        return {"text": text}

    @staticmethod
    def extract_text_from_pdf(input_data: Union[str, bytes]) -> str:
        """
        Extracts text from a PDF file or bytes.
        """
        try:
            doc = None
            if isinstance(input_data, str):
                if not os.path.exists(input_data):
                    logger.error(f"File not found: {input_data}")
                    raise FileNotFoundError(f"File not found: {input_data}")
                logger.debug(f"Opening PDF from path: {input_data}")
                doc = fitz.open(input_data)
            
            elif isinstance(input_data, bytes):
                logger.debug("Opening PDF from bytes")
                doc = fitz.open(stream=input_data, filetype="pdf")
            
            else:
                raise ValueError("Input must be a file path (str) or bytes.")

            text = ""
            for page in doc:
                text += page.get_text()
            
            logger.info(f"Successfully extracted {len(text)} characters from PDF.")
            return text.strip()

        except Exception as e:
            source = input_data if isinstance(input_data, str) else "bytes"
            logger.error(f"Failed to process PDF {source}: {str(e)}", exc_info=True)
            raise Exception(f"Failed to process PDF {source}: {str(e)}")

    @staticmethod
    def extract_text_from_docx(input_data: Union[str, bytes]) -> str:
        """
        Extracts text from a DOCX file or bytes.
        """
        try:
            doc = None
            if isinstance(input_data, str):
                if not os.path.exists(input_data):
                    raise FileNotFoundError(f"File not found: {input_data}")
                logger.debug(f"Opening DOCX from path: {input_data}")
                doc = docx.Document(input_data)
            
            elif isinstance(input_data, bytes):
                logger.debug("Opening DOCX from bytes")
                # python-docx can open file-like objects
                file_stream = io.BytesIO(input_data)
                doc = docx.Document(file_stream)
            
            else:
                raise ValueError("Input must be a file path (str) or bytes.")

            text = []
            for para in doc.paragraphs:
                text.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            text.append(para.text)

            full_text = "\n".join(text)
            logger.info(f"Successfully extracted {len(full_text)} characters from DOCX.")
            return full_text.strip()

        except Exception as e:
            source = input_data if isinstance(input_data, str) else "bytes"
            logger.error(f"Failed to process DOCX {source}: {str(e)}", exc_info=True)
            raise Exception(f"Failed to process DOCX {source}: {str(e)}")
