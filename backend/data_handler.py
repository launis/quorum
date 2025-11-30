import io
import docx
import PyPDF2
import re
from typing import Tuple, Dict, Any

class TextUpload:
    """Helper class to treat text as an uploaded file."""
    def __init__(self, text, name="pasted_text.txt"):
        self.text = text
        self.name = name
        self.type = "text/plain"
        self._value = text.encode('utf-8')

    def getvalue(self):
        return self._value

class DataHandler:
    """
    Data Layer: Handles file reading and processing.
    Adapted from 'HolistinenKoodit' repository.
    """
    
    def read_file_content(self, uploaded_file: Any) -> str:
        """Reads file content as text."""
        # Determine file type and content
        file_type = ""
        file_obj = None

        # Check for FastAPI UploadFile
        if hasattr(uploaded_file, 'content_type'):
            file_type = uploaded_file.content_type
            file_obj = uploaded_file.file
        # Check for Streamlit UploadedFile
        elif hasattr(uploaded_file, 'type'):
            file_type = uploaded_file.type
            file_obj = uploaded_file
        
        if file_type == "application/pdf":
            return self._read_pdf(file_obj)
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
             return self._read_docx(file_obj)
        else:
            # Assume text file
            try:
                # FastAPI
                if hasattr(uploaded_file, 'read'):
                    # Reset cursor just in case
                    uploaded_file.file.seek(0)
                    content = uploaded_file.file.read()
                # Streamlit
                elif hasattr(uploaded_file, 'getvalue'):
                    content = uploaded_file.getvalue()
                else:
                    return str(uploaded_file)

                if isinstance(content, bytes):
                    return content.decode("utf-8")
                return str(content)
            except Exception:
                return str(uploaded_file)

    def _read_pdf(self, file_obj) -> str:
        """Reads text from PDF using pdfplumber (robust)."""
        try:
            import pdfplumber
            text = ""
            with pdfplumber.open(file_obj) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text
        except ImportError:
             # Fallback to PyPDF2 if pdfplumber is not installed
             try:
                reader = PyPDF2.PdfReader(file_obj)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
             except Exception as e:
                 return f"Error reading PDF file (PyPDF2): {str(e)}"
        except Exception as e:
            return f"Error reading PDF file: {str(e)}"

    def _read_docx(self, file_path_or_obj) -> str:
        """Reads text from DOCX (including tables)."""
        try:
            doc = docx.Document(file_path_or_obj)
            text = []
            for element in doc.element.body:
                if element.tag.endswith('p'): # Paragraph
                    para = docx.text.paragraph.Paragraph(element, doc)
                    text.append(para.text)
                elif element.tag.endswith('tbl'): # Table
                    table = docx.table.Table(element, doc)
                    for row in table.rows:
                        row_text = [cell.text for cell in row.cells]
                        text.append(" | ".join(row_text))
            return "\n".join(text)
        except Exception as e:
            return f"Error reading DOCX file: {str(e)}"

    def parse_prompt_modules(self, file_path_or_obj) -> Tuple[str, Dict[str, str]]:
        """
        Parses the Master Assessment Prompt into parts:
        1. Common Rules (text before STEP 1)
        2. Steps (STEP 1, STEP 2, etc.)
        
        Returns: (common_rules, phases_dict)
        """
        try:
            doc = docx.Document(file_path_or_obj)
            common_rules = []
            phases = {}
            current_phase = None
            
            # Regex to identify steps: "VAIHE" + spaces + number
            # Adapted to support "STEP" as well if needed, but sticking to Finnish "VAIHE" as per source
            phase_pattern = re.compile(r"^(VAIHE\s+\d+|STEP\s+\d+)", re.IGNORECASE)
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                    
                # Identify phase change
                match = phase_pattern.match(text)
                if match:
                    # Normalize key: "VAIHE 1" (always uppercase, single space)
                    phase_key = match.group(1).upper().replace("  ", " ")
                    
                    current_phase = phase_key
                    phases[current_phase] = [text] 
                    continue

                if current_phase:
                    phases[current_phase].append(text)
                else:
                    common_rules.append(text)
            
            # Convert lists to text
            common_text = "\n".join(common_rules)
            phases_text = {k: "\n".join(v) for k, v in phases.items()}
            
            return common_text, phases_text
            
        except Exception as e:
            return f"Error parsing prompt: {str(e)}", {}
