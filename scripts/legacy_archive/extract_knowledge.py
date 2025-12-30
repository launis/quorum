import docx
import json
import re
import os
import logging

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')
logger = logging.getLogger(__name__)

DOC_PATH = os.path.join("data", "Holistinen Mestaruus.docx")
OUTPUT_PATH = os.path.join("data", "coach_resources.json")

def clean_text(text):
    return re.sub(r'\s+', ' ', text).strip()

def extract_knowledge():
    if not os.path.exists(DOC_PATH):
        logger.error(f"File not found: {DOC_PATH}")
        return

    doc = docx.Document(DOC_PATH)
    logger.info(f"Opened document with {len(doc.paragraphs)} paragraphs.")

    data = {
        "concepts": {
            "Bloom": [],
            "Toulmin": [],
            "BARS": []
        },
        "citations": [],
        "bibliography": []
    }

    current_concept = None
    
    # 1. Main Content Parsing (Heuristic)
    for para in doc.paragraphs:
        text = clean_text(para.text)
        if not text: continue
        
        text_lower = text.lower()
        
        # Detect headers/topics
        if "bloom" in text_lower:
            current_concept = "Bloom"
        elif "toulmin" in text_lower:
            current_concept = "Toulmin"
        elif "bars" in text_lower or "behaviorally anchored" in text_lower:
            current_concept = "BARS"
        elif "hybridirubriikki" in text_lower or "hybridirubriikin" in text_lower:
             current_concept = "Hybrid Rubric"
        elif "adversariaali" in text_lower or "debatti" in text_lower or "väittely" in text_lower:
             current_concept = "Adversarial Debate"
        elif "performatiivi" in text_lower or "goodhart" in text_lower:
             current_concept = "Performative Reflection"
        elif "automaatioharha" in text_lower or "automation bias" in text_lower:
             current_concept = "Automation Bias"
        elif "lähdeluettelo" in text_lower or "references" in text_lower:
            current_concept = "Bibliography"
            
        # Store content
        if current_concept == "Bibliography":
             # Try to capture full bibliography entries
             if len(text) > 20: 
                 data["bibliography"].append(text)
        elif current_concept:
            if current_concept not in data["concepts"]:
                data["concepts"][current_concept] = []
            data["concepts"][current_concept].append(text)

        # 2. Inline Citation Extraction
        # Look for patterns like (Name 2023) or (Name et al. 2023)
        citations = re.findall(r'\((?:[A-Z][a-z]+(?: ym\.| et al\.)?,? \d{4}(?:, \d{1,4})?)\)', text)
        if citations:
            for cit in citations:
                if cit not in data["citations"]:
                    data["citations"].append(cit)

    # 3. Refine Output
    final_output = {
        "concepts": {},
        "references": {
            "inline_citations": sorted(list(set(data["citations"]))),
            "bibliography": data["bibliography"][-15:] # Take last 15 entries as likely biblio, rudimentary
        }
    }
    
    for concept, paragraphs in data["concepts"].items():
        # Join paragraphs into a single text block
        full_text = " ".join(paragraphs)
        # Limit length to avoid massive context dumping? No, let's keep rich context.
        final_output["concepts"][concept] = full_text

    # Write to JSON
    with open(OUTPUT_PATH, 'w', encoding='utf-8') as f:
        json.dump(final_output, f, ensure_ascii=False, indent=2)
        
    logger.info(f"Extraction complete. Saved to {OUTPUT_PATH}")
    logger.info(f"Found {len(final_output['references']['inline_citations'])} inline citations and {len(final_output['references']['bibliography'])} bibliography entries.")

if __name__ == "__main__":
    extract_knowledge()
