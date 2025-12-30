import sys
import os
import io
import docx
from unittest.mock import MagicMock

# Add project root to path
sys.path.append(os.getcwd())

from backend.services.knowledge_base_service import KnowledgeBaseService
from backend.services.knowledge_base_parser import KnowledgeBaseParser

def test_kb_ingestion():
    print("Testing KB Ingestion...")
    
    # 1. Mock Repository
    repo = MagicMock()
    
    # 2. Mock Storage
    storage = MagicMock()
    storage.save = MagicMock(return_value="stored_path/test.docx")
    
    # 3. Create Service
    service = KnowledgeBaseService(repo, storage)
    
    # 4. Create Fake DOCX Bytes
    doc = docx.Document()
    doc.add_heading("Concept Alpha", level=1)
    doc.add_paragraph("Definition of alpha.")
    doc.add_heading("Lähdeluettelo", level=1)
    doc.add_paragraph("Citation 1 (10.1000/xyz)")
    
    byte_io = io.BytesIO()
    doc.save(byte_io)
    file_bytes = byte_io.getvalue()
    
    # 5. Run Ingest
    res = service.ingest_from_bytes(file_bytes, "test.docx")
    
    # 6. Verify
    print("Result:", res)
    
    # Verify Storage Called
    storage.save.assert_called_once()
    print("Storage.save called:", storage.save.called)
    
    # Verify Repo Called
    print("Repo.add_knowledge_base_item call count:", repo.add_knowledge_base_item.call_count)
    if repo.add_knowledge_base_item.call_count >= 2:
        print("SUCCESS: Items added to repository.")
    else:
        print("FAILURE: Items NOT added to repository.")

if __name__ == "__main__":
    test_kb_ingestion()
